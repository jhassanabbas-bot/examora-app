# app.py — Examora (Beta) Streamlit MVP
# ============================================================
# CHANGELOG
# 1. Simplified UI flow — auto doc-type detection (threshold: 300 pages)
# 2. TOC / chapter auto-detection from PDF outline
# 3. Bookmark / resume per user per PDF
# 4. Better single-doc handling — coverage stats, scanned PDF warning
# 5. Speed optimisations — gpt-4o-mini, smaller input budgets, dynamic max_tokens
# 6. Professional in-page UI — gradient headers, styled cards, score banners
# 7. GA4 client-side events (Streamlit-safe, deduplicated)
# 8. Adaptive summarisation — short/medium/long map-reduce
# 9. WORD DOCUMENT SUPPORT (this update)
#    - Accepts .pdf and .docx uploads
#    - extract_text_from_docx() pulls full text from all paragraphs and tables
#    - Estimates page count from word count (~250 words/page) for .docx
#    - Word files always use Single Document Mode (no chapter splitting)
#    - Book Mode only available for PDF uploads
#    - All downstream summarisation and exam generation unchanged
# ============================================================

import os
import re
import json
import time
import random
import io
import csv
import secrets
import hashlib
import uuid
from pathlib import Path
from datetime import datetime
from urllib.parse import quote_plus

import requests
import streamlit as st
import streamlit.components.v1 as components
from pypdf import PdfReader
from docx import Document as DocxDocument
from dotenv import load_dotenv
from openai import OpenAI
from openai import AuthenticationError, RateLimitError, BadRequestError


# ============================================================
# Streamlit config — MUST be first Streamlit call
# ============================================================
st.set_page_config(page_title="Examora (Beta)", layout="wide")


# ============================================================
# ENV
# ============================================================
load_dotenv(dotenv_path=r".\.env", override=True)

API_KEY        = os.getenv("OPENAI_API_KEY") or ""
BETA_LIMIT     = int(os.getenv("EXAMORA_BETA_LIMIT") or "20")
ADMIN_EMAIL    = (os.getenv("EXAMORA_ADMIN_EMAIL") or "").strip().lower()
ADMIN_PASSWORD = os.getenv("EXAMORA_ADMIN_PASSWORD") or ""
RESET_CODE     = (os.getenv("EXAMORA_RESET_CODE") or "").strip()

GA_MEASUREMENT_ID = (os.getenv("GA4_MEASUREMENT_ID") or "G-GGZNKBCS1E").strip()
GA_API_SECRET     = (os.getenv("GA4_API_SECRET") or "").strip()

FEEDBACK_URL = (
    "https://docs.google.com/forms/d/e/"
    "1FAIpQLSc1n_uvwsnr1NpXiY_1SCg5_t_6MnsWVgG54z2NZHgVJOrkVw/viewform?usp=header"
)

# ── Thresholds ──────────────────────────────────────────────
# Docs below this page count are treated as Single Documents.
# Raised from 80 → 300 so AAPM TG reports (up to ~200 pages) stay
# in Single Document Mode where they belong.
BOOK_PAGE_THRESHOLD = 300

# Adaptive summarisation character thresholds
# Short  : < 40,000 chars  (~50 pages)   → 1 API call
# Medium : 40k – 120k chars (~50–150 pg) → 3-chunk map-reduce
# Long   : > 120,000 chars  (~150+ pg)   → 5-chunk map-reduce + overview
SUMMARY_SHORT_LIMIT  =  40_000
SUMMARY_MEDIUM_LIMIT = 120_000

client = OpenAI()


# ============================================================
# GA4 — client-side events (Streamlit-safe)
# ============================================================
def ga_init() -> None:
    if not GA_MEASUREMENT_ID or st.session_state.get("_ga_loaded"):
        return
    st.session_state["_ga_loaded"] = True
    st.session_state.setdefault("_ga_sent", set())
    nonce = f"{int(time.time()*1000)}_{random.randint(1,10_000_000)}"
    components.html(f"""
        <!-- ga_init_nonce:{nonce} -->
        <script async src="https://www.googletagmanager.com/gtag/js?id={GA_MEASUREMENT_ID}"></script>
        <script>
          window.dataLayer = window.dataLayer || [];
          function gtag(){{dataLayer.push(arguments);}}
          window.gtag = window.gtag || gtag;
          gtag('js', new Date());
          gtag('config', '{GA_MEASUREMENT_ID}', {{'send_page_view': false}});
        </script>""", height=0)

def ga_event(name: str, params: dict | None = None, once_key: str | None = None) -> None:
    if not GA_MEASUREMENT_ID:
        return
    ga_init()
    if once_key:
        sent = st.session_state.setdefault("_ga_sent", set())
        if once_key in sent:
            return
        sent.add(once_key)
    payload = json.dumps(params or {}, ensure_ascii=False).replace("</", "<\\/")
    nonce   = f"{name}_{abs(hash(payload))}_{int(time.time()*1000)}_{random.randint(1,10_000)}"
    components.html(f"""
        <!-- ga_evt_nonce:{nonce} -->
        <script>
          (function() {{
            const payload = {payload};
            function trySend() {{
              try {{ if (typeof window.gtag === "function") {{ window.gtag("event", "{name}", payload); return true; }} }} catch(e) {{}}
              return false;
            }}
            if (trySend()) return;
            let attempts = 0;
            const t = setInterval(() => {{ attempts++; if (trySend() || attempts >= 40) clearInterval(t); }}, 50);
          }})();
        </script>""", height=0)


# ============================================================
# Local data store
# ============================================================
DATA_DIR = Path(".examora")
DATA_DIR.mkdir(exist_ok=True)

USAGE_FILE     = DATA_DIR / "beta_usage.json"
USERS_FILE     = DATA_DIR / "users.json"
FEEDBACK_LOG   = DATA_DIR / "feedback_log.json"
BOOKMARKS_FILE = DATA_DIR / "bookmarks.json"


# ============================================================
# Utilities
# ============================================================
def now_ts() -> int:
    return int(time.time())

def ts_to_str(ts: int) -> str:
    try:
        return datetime.fromtimestamp(int(ts)).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return ""

def normalize_email(e: str) -> str:
    return (e or "").strip().lower()

def looks_like_email(e: str) -> bool:
    return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", normalize_email(e)))

def _seek0(f):
    try:
        f.seek(0)
    except Exception:
        pass

def build_prefilled_feedback_url(base: str, email: str) -> str:
    eid = (os.getenv("EXAMORA_FEEDBACK_EMAIL_ENTRY_ID") or "").strip()
    if not email or not looks_like_email(email) or not eid:
        return base
    sep = "&" if "?" in base else "?"
    return f"{base}{sep}usp=pp_url&entry.{eid}={quote_plus(email)}"


# ============================================================
# Safe file IO
# ============================================================
def _read_json(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default

def _write_json(path: Path, data):
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


# ============================================================
# Feedback log
# ============================================================
def load_feedback_log() -> list:
    d = _read_json(FEEDBACK_LOG, [])
    return d if isinstance(d, list) else []

def append_feedback_log(item: dict) -> None:
    log = load_feedback_log()
    log.append(item)
    _write_json(FEEDBACK_LOG, log[-300:])


# ============================================================
# Cached IO
# ============================================================
@st.cache_data(ttl=30, show_spinner=False)
def load_usage() -> dict:
    return _read_json(USAGE_FILE, {})

@st.cache_data(ttl=60, show_spinner=False)
def load_users() -> dict:
    return _read_json(USERS_FILE, {})

def save_usage(data: dict) -> None:
    _write_json(USAGE_FILE, data)
    load_usage.clear()

def save_users(users: dict) -> None:
    _write_json(USERS_FILE, users)
    load_users.clear()


# ============================================================
# Bookmark store
# ============================================================
def _bm_key(email: str, pdf: str) -> str:
    return f"{normalize_email(email)}||{pdf}"

def load_bookmark(email: str, pdf: str) -> int | None:
    bm = _read_json(BOOKMARKS_FILE, {})
    v  = bm.get(_bm_key(email, pdf))
    return int(v) if v is not None else None

def save_bookmark(email: str, pdf: str, idx: int) -> None:
    bm = _read_json(BOOKMARKS_FILE, {})
    bm[_bm_key(email, pdf)] = idx
    _write_json(BOOKMARKS_FILE, bm)

def clear_bookmark(email: str, pdf: str) -> None:
    bm = _read_json(BOOKMARKS_FILE, {})
    bm.pop(_bm_key(email, pdf), None)
    _write_json(BOOKMARKS_FILE, bm)


# ============================================================
# Usage store
# ============================================================
def get_user_usage(email: str) -> dict:
    return load_usage().get(
        normalize_email(email),
        {"exam_sessions_used": 0, "created_at": None, "last_used_at": None, "history": []},
    )

def increment_exam_session(email: str, meta: dict) -> int:
    email = normalize_email(email)
    usage = load_usage()
    if email not in usage:
        usage[email] = {"exam_sessions_used": 0, "created_at": now_ts(),
                        "last_used_at": None, "history": []}
    usage[email]["exam_sessions_used"] = int(usage[email].get("exam_sessions_used", 0)) + 1
    usage[email]["last_used_at"] = now_ts()
    hist = usage[email].get("history", [])
    hist.append(meta)
    usage[email]["history"] = hist[-50:]
    save_usage(usage)
    return int(usage[email]["exam_sessions_used"])

def usage_to_csv_bytes(usage: dict) -> bytes:
    out = io.StringIO()
    w   = csv.writer(out)
    w.writerow(["email", "exam_sessions_used", "created_at", "last_used_at", "history_count"])
    for em, row in usage.items():
        w.writerow([em, row.get("exam_sessions_used", 0),
                    ts_to_str(row.get("created_at") or 0),
                    ts_to_str(row.get("last_used_at") or 0),
                    len(row.get("history", []) or [])])
    return out.getvalue().encode("utf-8")


# ============================================================
# Users / auth
# ============================================================
PBKDF2_ITERS = 150_000

def pbkdf2_hash(pw: str, salt_hex: str) -> str:
    salt = bytes.fromhex(salt_hex)
    return hashlib.pbkdf2_hmac("sha256", pw.encode(), salt, PBKDF2_ITERS).hex()

def new_salt_hex() -> str:
    return secrets.token_bytes(16).hex()

def ensure_admin_user_exists() -> None:
    if st.session_state.get("_admin_checked"):
        return
    st.session_state["_admin_checked"] = True
    if not ADMIN_EMAIL or not ADMIN_PASSWORD:
        return
    users = load_users()
    if ADMIN_EMAIL in users:
        return
    salt = new_salt_hex()
    users[ADMIN_EMAIL] = {
        "salt": salt, "pwd_hash": pbkdf2_hash(ADMIN_PASSWORD, salt),
        "created_at": now_ts(), "last_login_at": None, "role": "admin",
    }
    save_users(users)

def register_user(email: str, pw: str) -> tuple[bool, str]:
    email = normalize_email(email)
    if not looks_like_email(email):
        return False, "Please enter a valid email."
    if len(pw) < 8:
        return False, "Password must be at least 8 characters."
    users = load_users()
    if email in users:
        return False, "Email already registered. Please log in."
    salt = new_salt_hex()
    users[email] = {
        "salt": salt, "pwd_hash": pbkdf2_hash(pw, salt),
        "created_at": now_ts(), "last_login_at": None, "role": "user",
    }
    save_users(users)
    return True, "Account created. You can log in now."

def verify_login(email: str, pw: str) -> tuple[bool, str]:
    email = normalize_email(email)
    users = load_users()
    row   = users.get(email)
    if not row:
        return False, "No account found. Please register."
    if pbkdf2_hash(pw, row.get("salt", "")) != row.get("pwd_hash", ""):
        return False, "Incorrect password."
    row["last_login_at"] = now_ts()
    users[email] = row
    save_users(users)
    return True, "Logged in."

def reset_password_with_code(email: str, code: str, new_pw: str) -> tuple[bool, str]:
    email = normalize_email(email)
    if not looks_like_email(email):
        return False, "Please enter a valid email."
    if not RESET_CODE:
        return False, "Reset not configured. Set EXAMORA_RESET_CODE in .env."
    if (code or "").strip() != RESET_CODE:
        return False, "Invalid reset code."
    if len(new_pw) < 8:
        return False, "New password must be at least 8 characters."
    users = load_users()
    if email not in users:
        return False, "No account found. Please register."
    salt = new_salt_hex()
    users[email]["salt"]         = salt
    users[email]["pwd_hash"]     = pbkdf2_hash(new_pw, salt)
    users[email]["pwd_reset_at"] = now_ts()
    save_users(users)
    return True, "Password reset. You can log in now."

def get_role(email: str) -> str:
    return (load_users().get(normalize_email(email), {}).get("role") or "user").lower()


# ============================================================
# OpenAI helpers
# ============================================================
def require_api_key():
    if not API_KEY:
        st.error("OPENAI_API_KEY not found. Check your .env file.")
        st.stop()

def safe_call(fn, *args, **kwargs):
    try:
        return fn(*args, **kwargs)
    except AuthenticationError:
        st.error("Authentication failed. Re-check OPENAI_API_KEY.")
    except RateLimitError as e:
        st.error(f"Rate limit: {e}")
    except BadRequestError as e:
        st.error(f"Bad request: {e}")
    except Exception as e:
        st.error(f"Unexpected error: {e}")
    return None


# ============================================================
# PDF extraction
# ============================================================
@st.cache_data(show_spinner=False)
def get_pdf_page_count(file_bytes: bytes) -> int:
    return len(PdfReader(io.BytesIO(file_bytes)).pages)

@st.cache_data(show_spinner=False)
def extract_text_cached(file_bytes: bytes, start_page: int,
                         end_page: int) -> tuple[str, int, int, int]:
    reader  = PdfReader(io.BytesIO(file_bytes))
    n_pages = len(reader.pages)
    start   = max(1, min(int(start_page), n_pages))
    end     = max(1, min(int(end_page),   n_pages))
    if end < start:
        start, end = end, start
    texts = [reader.pages[p].extract_text() or "" for p in range(start - 1, end)]
    return "\n".join(texts).strip(), n_pages, start, end

@st.cache_data(show_spinner=False)
def extract_toc_from_pdf(file_bytes: bytes, total_pages: int) -> list[dict]:
    reader   = PdfReader(io.BytesIO(file_bytes))
    chapters = []
    try:
        outline = reader.outline
        if outline:
            def _walk(items):
                for item in items:
                    if isinstance(item, list):
                        _walk(item)
                    else:
                        try:
                            pn = reader.get_destination_page_number(item) + 1
                            t  = (item.title or "").strip()
                            if t and pn:
                                chapters.append({"title": t, "start": pn})
                        except Exception:
                            pass
            _walk(outline)
    except Exception:
        pass
    if not chapters:
        pat = re.compile(
            r"^(chapter\s+\d+[\.\:]?\s*.{0,60}|section\s+\d+[\.\:]?\s*.{0,60})",
            re.IGNORECASE | re.MULTILINE,
        )
        for p in range(min(total_pages, 15)):
            try:
                txt = reader.pages[p].extract_text() or ""
                for m in pat.finditer(txt):
                    chapters.append({"title": m.group(0).strip()[:80], "start": p + 1})
            except Exception:
                pass
    if not chapters:
        return []
    seen, unique = set(), []
    for c in sorted(chapters, key=lambda x: x["start"]):
        k = (c["start"], c["title"][:30])
        if k not in seen:
            seen.add(k)
            unique.append(c)
    for i, c in enumerate(unique):
        c["end"] = unique[i + 1]["start"] - 1 if i + 1 < len(unique) else total_pages
        c["end"] = max(c["start"], c["end"])
    return unique


# ============================================================
# Word document (.docx) extraction
# ============================================================
def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """
    Extract all text from a .docx file.
    Returns (text, estimated_pages).
    Pages estimated at ~250 words per page — not exact but
    good enough for doc-size classification and UI display.
    Tables are included: each cell extracted as a line.
    """
    doc   = DocxDocument(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    text       = "\n".join(parts).strip()
    word_count = len(text.split())
    est_pages  = max(1, round(word_count / 250))
    return text, est_pages


# ============================================================
# Text helpers
# ============================================================
def chunk_text(text: str, max_chars: int = 8000) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    return [text[i: i + max_chars] for i in range(0, len(text), max_chars)]

def prepare_input(text: str, char_budget: int) -> str:
    chunks = chunk_text(text)
    if not chunks:
        return ""
    parts, total = [], 0
    for chunk in chunks:
        if total + len(chunk) > char_budget:
            break
        parts.append(chunk)
        total += len(chunk)
    return "\n\n".join(parts) if parts else chunks[0][:char_budget]

def classify_doc_size(text: str) -> str:
    """
    Classify document by character count for adaptive summarisation.
    short  : < 40,000 chars  (~50 pages)
    medium : 40k – 120k chars (~50–150 pages)
    long   : > 120,000 chars  (~150+ pages)
    """
    n = len(text)
    if n < SUMMARY_SHORT_LIMIT:
        return "short"
    if n < SUMMARY_MEDIUM_LIMIT:
        return "medium"
    return "long"

def text_coverage_info(text: str, sp: int, ep: int) -> str:
    chars = len(text)
    size  = classify_doc_size(text)
    size_labels = {"short": "Short document", "medium": "Medium document", "long": "Long document"}
    return (f"Pages {sp}–{ep} &nbsp;·&nbsp; {chars:,} chars "
            f"&nbsp;·&nbsp; ~{chars // 4:,} tokens "
            f"&nbsp;·&nbsp; {size_labels[size]}")

def _extract_json(text: str) -> dict:
    if not text:
        raise ValueError("Empty model output.")
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not m:
        raise ValueError("Could not locate JSON in model output.")
    return json.loads(m.group(0))

def _difficulty_hint(level: str) -> str:
    if level == "Easy":
        return "Easy: foundational definitions, basic recall, simple conceptual checks."
    if level == "Hard":
        return "Hard: multi-step reasoning, subtle distinctions, application questions."
    return "Medium: conceptual understanding and typical exam-style reasoning."

def get_text_hash(text: str) -> str:
    key, lk = "_text_hash_cache", "_text_hash_len"
    if st.session_state.get(lk) != len(text) or key not in st.session_state:
        st.session_state[key] = hashlib.md5(text.encode()).hexdigest()
        st.session_state[lk]  = len(text)
    return st.session_state[key]


# ============================================================
# LLM helpers — single chunk summary + combine
# ============================================================
def _llm(system: str, user: str, max_tokens: int = 900) -> str:
    """Single gpt-4o-mini call. Returns content string."""
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system},
            {"role": "user",   "content": user},
        ],
        max_tokens=max_tokens,
        temperature=0.3,
    )
    return resp.choices[0].message.content.strip()


def _summarise_chunk(chunk_text_: str, chunk_label: str) -> str:
    """
    Summarise one chunk of text into bullet points.
    Used as the MAP step in map-reduce.
    """
    system = (
        "You are Examora, an expert medical study tutor. "
        f"Summarise the following section of a document ({chunk_label}). "
        "Extract only the most important points as concise bullet points. "
        "Be faithful to the material. Do not invent details."
    )
    return _llm(system, chunk_text_, max_tokens=600)


def _combine_summaries(chunk_summaries: list[str], doc_size: str) -> str:
    """
    REDUCE step: combine chunk summaries into the final structured summary.
    Output structure adapts to document size.
    """
    combined = "\n\n---\n\n".join(
        f"[Section {i+1}]\n{s}" for i, s in enumerate(chunk_summaries)
    )

    if doc_size == "long":
        structure = (
            "## Document Overview\n"
            "Write 2–3 sentences summarising the overall scope and purpose of this document.\n\n"
            "## Key Concepts\n"
            "- 12–18 bullet points covering the most important concepts across all sections.\n\n"
            "## Common Confusions\n"
            "- 4–6 bullet points on concepts candidates typically mix up.\n\n"
            "## Exam Takeaways\n"
            "- 8 bullet points the candidate must remember for a board exam."
        )
        token_limit = 1200
    elif doc_size == "medium":
        structure = (
            "## Key Concepts\n"
            "- 10–14 bullet points covering the most important ideas.\n\n"
            "## Common Confusions\n"
            "- 3–5 bullet points on concepts candidates typically mix up.\n\n"
            "## Exam Takeaways\n"
            "- 6 bullet points the candidate must remember for a board exam."
        )
        token_limit = 1000
    else:
        structure = (
            "## Key Concepts\n"
            "- 8–12 bullet points covering the most important ideas.\n\n"
            "## Common Confusions\n"
            "- 3–5 bullet points on concepts candidates typically mix up.\n\n"
            "## Exam Takeaways\n"
            "- 5 bullet points the candidate must remember for a board exam."
        )
        token_limit = 900

    system = (
        "You are Examora, an expert medical study tutor. "
        "You have been given bullet-point summaries of different sections of a document. "
        "Synthesise them into ONE coherent study summary using exactly this structure:\n\n"
        f"{structure}\n\n"
        "Rules: be faithful to the material only. Do not invent details. "
        "Keep bullets concise (1–2 sentences max). Remove duplicates across sections."
    )
    return _llm(system, combined, max_tokens=token_limit)


# ============================================================
# Adaptive summarisation — cached by text hash
# ============================================================
@st.cache_data(show_spinner=False)
def summarize_text_cached(text_hash: str, text: str) -> str:
    """
    Adaptive map-reduce summarisation:
    - Short  (<40k chars):  1 API call, concise output
    - Medium (40k-120k):    3 chunks → map → reduce (2 calls total per chunk + 1 combine)
    - Long   (>120k chars): 5 chunks → map → reduce
    All results cached — second call on same text is instant.
    """
    doc_size = classify_doc_size(text)
    n_chars  = len(text)

    if doc_size == "short":
        # Single call — fast path
        input_text = prepare_input(text, char_budget=24_000)
        system = (
            "You are Examora, an expert medical study tutor. "
            "Create a concise study summary with exactly these three sections:\n\n"
            "## Key Concepts\n- 8–12 bullet points.\n\n"
            "## Common Confusions\n- 3–5 bullet points.\n\n"
            "## Exam Takeaways\n- 5 bullet points for a board exam candidate.\n\n"
            "Rules: faithful to material only. No invented details. "
            "1–2 sentences per bullet max."
        )
        return _llm(system, input_text, max_tokens=900)

    # Medium / Long: split into n_chunks equal-sized segments, summarise each, then combine
    n_chunks   = 3 if doc_size == "medium" else 5
    chunk_size = n_chars // n_chunks
    segments   = [text[i * chunk_size: (i + 1) * chunk_size] for i in range(n_chunks)]
    # Make sure the last segment captures any remainder
    segments[-1] = text[(n_chunks - 1) * chunk_size:]

    chunk_summaries = []
    for i, seg in enumerate(segments):
        label   = f"Part {i+1} of {n_chunks}"
        trimmed = seg[:24_000]   # cap each segment before sending
        chunk_summaries.append(_summarise_chunk(trimmed, label))

    return _combine_summaries(chunk_summaries, doc_size)


# ============================================================
# MCQ generation — cached
# ============================================================
@st.cache_data(show_spinner=False)
def generate_mcqs_cached(text_hash: str, n_questions: int,
                          difficulty: str, text: str) -> dict:
    input_text         = prepare_input(text, char_budget=28_000)
    if not input_text:
        return {"questions": []}
    diff_hint          = _difficulty_hint(difficulty)
    dynamic_max_tokens = min(n_questions * 120 + 300, 2500)
    system_prompt = (
        "You are Examora, an expert board exam question writer. "
        "Generate multiple-choice questions using ONLY the provided material — no hallucinations.\n\n"
        f"Difficulty: {diff_hint}\n\n"
        "Return ONLY a valid JSON object — no markdown, no preamble:\n"
        '{"questions":[{"q":"...","options":{"A":"...","B":"...","C":"...","D":"..."},'
        '"answer":"A|B|C|D","explanation":"one concise sentence"}]}\n\n'
        f"Generate exactly {n_questions} questions."
    )
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": input_text},
        ],
        max_tokens=dynamic_max_tokens,
        temperature=0.3,
    )
    data  = _extract_json(resp.choices[0].message.content)
    qs    = data.get("questions", [])
    clean = []
    for item in qs:
        if not isinstance(item, dict):
            continue
        qtext = (item.get("q") or "").strip()
        opts  = item.get("options") or {}
        ans   = item.get("answer")
        exp   = (item.get("explanation") or "").strip()
        if (qtext and isinstance(opts, dict)
                and all(k in opts for k in ["A","B","C","D"])
                and ans in ["A","B","C","D"]):
            clean.append({
                "q": qtext,
                "options": {k: str(opts[k]).strip() for k in ["A","B","C","D"]},
                "answer": ans, "explanation": exp,
            })
    return {"questions": clean}

def shuffle_question_options(q: dict) -> dict:
    labels       = ["A","B","C","D"]
    correct_text = q["options"][q["answer"]]
    texts        = [q["options"][l] for l in labels]
    random.shuffle(texts)
    new_opts    = {labels[i]: texts[i] for i in range(4)}
    new_correct = next((k for k, v in new_opts.items() if v == correct_text), None)
    return {**q, "options": new_opts, "answer": new_correct} if new_correct else q

def build_results_csv(questions: list, answers: dict) -> bytes:
    out = io.StringIO()
    w   = csv.writer(out)
    w.writerow(["Question#","Question","YourAnswer","CorrectAnswer","Correct?","Explanation"])
    for i, q in enumerate(questions, 1):
        user    = answers.get(i)
        correct = q["answer"]
        w.writerow([i, q["q"], user or "", correct,
                    "YES" if user == correct else "NO", q.get("explanation","")])
    return out.getvalue().encode("utf-8")


# ============================================================
# Session state
# ============================================================
def init_state():
    if st.session_state.get("_initialized"):
        return
    st.session_state["_initialized"] = True

    st.session_state.setdefault("auth_email",  "")
    st.session_state.setdefault("is_authed",   False)

    st.session_state.setdefault("summary_open", False)
    st.session_state.setdefault("summary_text", "")

    st.session_state.setdefault("exam_open",  False)
    st.session_state.setdefault("questions",  [])
    st.session_state.setdefault("answers",    {})
    st.session_state.setdefault("q_index",    0)
    st.session_state.setdefault("flagged",    set())
    st.session_state.setdefault("submitted",  False)

    st.session_state.setdefault("doc_mode",   None)
    st.session_state.setdefault("scope_mode", "")

    st.session_state.setdefault("book_chapter_idx",     0)
    st.session_state.setdefault("manual_chapters_text", "")
    st.session_state.setdefault("quick_pages_per",      20)
    st.session_state.setdefault("toc_chapters",         None)

    st.session_state.setdefault("uploaded_file_bytes", None)
    st.session_state.setdefault("uploaded_file_name",  None)

def reset_exam_state():
    st.session_state.exam_open = False
    st.session_state.questions = []
    st.session_state.answers   = {}
    st.session_state.q_index   = 0
    st.session_state.flagged   = set()
    st.session_state.submitted = False
    for k in list(st.session_state.keys()):
        if str(k).startswith("choice_"):
            del st.session_state[k]

def reset_doc_state():
    reset_exam_state()
    st.session_state.summary_open = False
    st.session_state.summary_text = ""
    st.session_state.toc_chapters = None
    st.session_state.book_chapter_idx = 0
    st.session_state.pop("_last_extracted_key", None)


init_state()
ensure_admin_user_exists()
ga_init()


# ============================================================
# Global CSS
# ============================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.block-container { padding-top: 1rem !important; padding-bottom: 3rem !important; }

.ex-hero {
  background: linear-gradient(135deg, #1a3a6b 0%, #2563eb 100%);
  border-radius: 16px; padding: 28px 36px 22px 36px;
  margin-bottom: 20px; color: #fff;
}
.ex-hero-badge {
  display:inline-block; background:rgba(255,255,255,0.18);
  border:1px solid rgba(255,255,255,0.3); border-radius:20px;
  padding:3px 12px; font-size:12px; font-weight:600; color:#fff; margin-bottom:10px;
}
.ex-hero-title { font-size:2rem; font-weight:900; color:#fff !important; margin:0 0 4px 0; letter-spacing:-0.5px; }
.ex-hero-sub   { color:rgba(255,255,255,0.78); font-size:15px; margin:0; }

.ex-section-label {
  font-size:11px; font-weight:700; letter-spacing:1px;
  text-transform:uppercase; color:#6b7280; margin:24px 0 8px 0;
}
.ex-status {
  background:#f0f7ff; border:1px solid #bfdbfe;
  border-left:4px solid #2563eb; border-radius:10px;
  padding:10px 16px; font-size:14px; margin-bottom:14px; color:#1e3a5f;
}
.ex-status-long {
  background:#fefce8; border:1px solid #fde68a;
  border-left:4px solid #f59e0b; border-radius:10px;
  padding:10px 16px; font-size:14px; margin-bottom:14px; color:#78350f;
}
.bookmark-banner {
  background:#fffbeb; border:1px solid #fcd34d;
  border-left:4px solid #f59e0b; border-radius:10px;
  padding:10px 16px; font-size:14px; margin-bottom:14px; color:#78350f;
}

/* Summary */
.sum-wrap { border:1px solid #e5e7eb; border-radius:14px; overflow:hidden; margin-bottom:20px; box-shadow:0 4px 20px rgba(0,0,0,0.07); }
.sum-header { background:linear-gradient(135deg,#1a3a6b,#2563eb); padding:18px 24px 14px 24px; }
.sum-header-title { font-size:1.15rem; font-weight:900; color:#fff !important; margin:0 0 2px 0; }
.sum-header-sub   { font-size:12px; color:rgba(255,255,255,0.75); margin:0; }
.sum-section { padding:16px 24px 8px 24px; border-bottom:1px solid #f3f4f6; background:#fff; }
.sum-section:last-child { border-bottom:none; }
.sum-section-title { font-size:13px; font-weight:800; color:#1a3a6b; margin-bottom:10px; }
.sum-footer { background:#f9fafb; padding:12px 24px; border-top:1px solid #f3f4f6; }

/* Exam */
.exam-wrap { border:1px solid #e5e7eb; border-radius:14px; overflow:hidden; margin-bottom:20px; box-shadow:0 4px 20px rgba(0,0,0,0.07); }
.exam-header { background:linear-gradient(135deg,#1a3a6b,#2563eb); padding:18px 24px 14px 24px; }
.exam-header-title { font-size:1.15rem; font-weight:900; color:#fff !important; margin:0 0 2px 0; }
.exam-header-sub   { font-size:12px; color:rgba(255,255,255,0.75); margin:0; }
.q-card { background:#f8fafc; border:1px solid #e2e8f0; border-radius:12px; padding:18px 22px 14px 22px; margin-bottom:14px; }
.q-meta { font-size:11px; font-weight:700; color:#2563eb; text-transform:uppercase; letter-spacing:0.8px; margin-bottom:8px; }
.q-text { font-size:16px; font-weight:600; color:#111827; line-height:1.55; }

/* Results */
.score-pass { background:linear-gradient(135deg,#065f46,#10b981); color:#fff; border-radius:12px; padding:22px 28px; text-align:center; margin-bottom:18px; }
.score-warn { background:linear-gradient(135deg,#92400e,#f59e0b); color:#fff; border-radius:12px; padding:22px 28px; text-align:center; margin-bottom:18px; }
.score-fail { background:linear-gradient(135deg,#7f1d1d,#ef4444); color:#fff; border-radius:12px; padding:22px 28px; text-align:center; margin-bottom:18px; }
.score-pct  { font-size:3.5rem; font-weight:900; line-height:1; }
.score-msg  { font-size:14px; opacity:0.88; margin-top:6px; }
.r-correct   { background:#f0fdf4; border:1px solid #86efac; border-radius:10px; padding:14px 18px; margin-bottom:8px; }
.r-incorrect { background:#fef2f2; border:1px solid #fca5a5; border-radius:10px; padding:14px 18px; margin-bottom:8px; }
.r-q   { font-size:14px; font-weight:600; color:#111827; margin-bottom:5px; }
.r-ans { font-size:13px; color:#374151; margin-bottom:3px; }
.r-exp { font-size:12px; color:#6b7280; margin-top:6px; font-style:italic; }

/* Buttons */
div.stButton > button { border-radius:9px !important; font-weight:700 !important; font-size:14px !important; }
div.stButton > button[kind="primary"] {
  background:linear-gradient(135deg,#1a3a6b,#2563eb) !important; border:none !important; color:#fff !important;
}
</style>
""", unsafe_allow_html=True)


# ============================================================
# Sidebar — Login / Register / Dashboard
# ============================================================
st.sidebar.markdown("## ⚕️ Examora")
st.sidebar.caption("Board exam preparation platform")

try:
    this_file = os.path.abspath(__file__)
    mtime = datetime.fromtimestamp(os.path.getmtime(this_file)).strftime("%Y-%m-%d %H:%M")
    st.sidebar.caption(f"v{mtime}")
except Exception:
    pass

auth_tab = st.sidebar.radio("Account", ["Login","Register","Dashboard"], index=0)

if auth_tab == "Login":
    email_in = st.sidebar.text_input("Email", value=st.session_state.auth_email,
                                     placeholder="name@example.com")
    pw_in    = st.sidebar.text_input("Password", type="password")
    if st.sidebar.button("Log in", use_container_width=True):
        ok, msg = verify_login(email_in, pw_in)
        if ok:
            st.session_state.is_authed  = True
            st.session_state.auth_email = normalize_email(email_in)
            st.rerun()
        else:
            st.sidebar.error(msg)
    with st.sidebar.expander("Forgot password?"):
        fp_email = st.text_input("Account email", key="fp_email", placeholder="name@example.com")
        fp_code  = st.text_input("Reset code",    key="fp_code",  type="password")
        fp_new1  = st.text_input("New password",  key="fp_new1",  type="password")
        fp_new2  = st.text_input("Confirm",        key="fp_new2",  type="password")
        if st.button("Reset password"):
            if fp_new1 != fp_new2:
                st.error("Passwords do not match.")
            else:
                ok, msg = reset_password_with_code(fp_email, fp_code, fp_new1)
                st.success(msg) if ok else st.error(msg)
        if not RESET_CODE:
            st.warning("Admin: set EXAMORA_RESET_CODE in .env.")
    st.sidebar.link_button("💬 Submit Feedback",
        build_prefilled_feedback_url(FEEDBACK_URL, normalize_email(email_in)))

elif auth_tab == "Register":
    reg_email = st.sidebar.text_input("Email", placeholder="name@example.com")
    reg_pw1   = st.sidebar.text_input("Password (min 8 chars)", type="password")
    reg_pw2   = st.sidebar.text_input("Confirm password",       type="password")
    if st.sidebar.button("Create account", use_container_width=True):
        if reg_pw1 != reg_pw2:
            st.sidebar.error("Passwords do not match.")
        else:
            ok, msg = register_user(reg_email, reg_pw1)
            st.sidebar.success(msg) if ok else st.sidebar.error(msg)
    st.sidebar.link_button("💬 Submit Feedback",
        build_prefilled_feedback_url(FEEDBACK_URL, normalize_email(reg_email)))

elif auth_tab == "Dashboard":
    if not st.session_state.is_authed:
        st.sidebar.warning("Please log in to view your dashboard.")
    else:
        de = st.session_state.auth_email
        dr = get_role(de)
        du = get_user_usage(de)
        st.sidebar.markdown(f"**{de}**")
        st.sidebar.markdown(f"Role: `{dr}`")
        used = du.get("exam_sessions_used", 0)
        st.sidebar.progress(used / BETA_LIMIT, text=f"{used}/{BETA_LIMIT} sessions used")
        if st.sidebar.button("Log out", use_container_width=True, key="sidebar_logout_dash"):
            st.session_state.is_authed           = False
            st.session_state.auth_email          = ""
            st.session_state.uploaded_file_bytes = None
            st.session_state.uploaded_file_name  = None
            reset_doc_state()
            st.rerun()
        st.sidebar.link_button("💬 Submit Feedback",
            build_prefilled_feedback_url(FEEDBACK_URL, de))
        if dr == "admin":
            st.sidebar.divider()
            st.sidebar.markdown("### Admin")
            st.sidebar.download_button(
                "⬇️ Download Usage CSV",
                data=usage_to_csv_bytes(load_usage()),
                file_name="examora_usage.csv", mime="text/csv",
            )
            fb = load_feedback_log()
            st.sidebar.caption(f"Feedback entries: {len(fb)}")
            for row in fb[-5:][::-1]:
                st.sidebar.write(f"- {ts_to_str(row.get('ts',0))} · {row.get('note','')[:40]}")


# ============================================================
# Hero header
# ============================================================
user_authed = st.session_state.is_authed
user_email  = st.session_state.auth_email if user_authed else ""
role        = get_role(user_email) if user_authed else "user"

st.markdown("""
<div class="ex-hero">
  <div class="ex-hero-badge">Beta</div>
  <div class="ex-hero-title">⚕️ Examora</div>
  <div class="ex-hero-sub">Professional board exam preparation &nbsp;·&nbsp; Upload · Summarise · Practise · Master</div>
</div>
""", unsafe_allow_html=True)

if not user_authed:
    st.info("👈 Please **Login** in the sidebar to use Examora.")
    st.stop()

st.sidebar.divider()
if st.sidebar.button("Log out", key="main_logout", use_container_width=True):
    st.session_state.is_authed           = False
    st.session_state.auth_email          = ""
    st.session_state.uploaded_file_bytes = None
    st.session_state.uploaded_file_name  = None
    reset_doc_state()
    st.rerun()

# Dashboard page
if auth_tab == "Dashboard":
    u    = get_user_usage(user_email)
    hist = u.get("history", []) or []
    st.markdown('<div class="ex-section-label">Your Dashboard</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    c1.metric("Sessions used",  f"{u.get('exam_sessions_used',0)} / {BETA_LIMIT}")
    c2.metric("First seen",     ts_to_str(u.get("created_at") or 0) if u.get("created_at") else "—")
    c3.metric("Last active",    ts_to_str(u.get("last_used_at") or 0) if u.get("last_used_at") else "—")
    if hist:
        st.markdown('<div class="ex-section-label">Recent Sessions</div>', unsafe_allow_html=True)
        for item in list(hist)[-10:][::-1]:
            stype = item.get("type","exam")
            tail  = (f"· {item.get('n_questions')} Qs · {item.get('difficulty')}"
                     if stype == "exam" else "")
            st.write(f"- {ts_to_str(item.get('ts'))} · **{stype}** · pages {item.get('pages')} {tail}")
    else:
        st.info("No session history yet.")
    st.stop()


# ============================================================
# STEP 1 — Upload document (PDF or Word)
# ============================================================
st.markdown('<div class="ex-section-label">Step 1 — Upload your document</div>', unsafe_allow_html=True)

# File type hint
st.caption("Accepted formats: PDF (.pdf) · Word document (.docx)")
uploaded = st.file_uploader(
    "Upload a PDF or Word document",
    type=["pdf", "docx"],
    label_visibility="collapsed",
)

if not uploaded:
    st.session_state.uploaded_file_bytes = None
    st.session_state.uploaded_file_name  = None
    st.session_state.doc_mode            = None
    st.info("📂 Upload a PDF or Word document (.docx) to begin.")
    st.stop()

if (st.session_state.uploaded_file_bytes is None
        or st.session_state.uploaded_file_name != uploaded.name):
    _seek0(uploaded)
    st.session_state.uploaded_file_bytes = uploaded.read()
    st.session_state.uploaded_file_name  = uploaded.name
    st.session_state.doc_mode            = None
    reset_doc_state()
    ga_event("upload_clicked",
             params={"file_name": uploaded.name},
             once_key=f"upload_{uploaded.name}")

file_bytes = st.session_state.uploaded_file_bytes
file_name  = uploaded.name
is_docx    = file_name.lower().endswith(".docx")

# ── For Word files: extract text immediately and lock to Single Doc mode ──
if is_docx:
    if st.session_state.get("_last_extracted_key") != ("docx", file_name, len(file_bytes)):
        with st.spinner("Reading Word document..."):
            docx_text, docx_est_pages = extract_text_from_docx(file_bytes)
        st.session_state["_docx_text"]      = docx_text
        st.session_state["_docx_est_pages"] = docx_est_pages
        st.session_state["_last_extracted_key"] = ("docx", file_name, len(file_bytes))
    else:
        docx_text      = st.session_state["_docx_text"]
        docx_est_pages = st.session_state["_docx_est_pages"]

    if len(docx_text) < 200:
        st.error(
            "⚠️ Very little text extracted from this Word file. "
            "It may be empty, image-only, or corrupted.")
        st.stop()

    # Word files always → Single Document Mode, no mode picker needed
    st.session_state.doc_mode  = "single"
    st.session_state.scope_mode = "Single Document (Word)"
    total_pages = docx_est_pages
    pdf_name    = file_name

    doc_size   = classify_doc_size(docx_text)
    status_cls = "ex-status-long" if doc_size == "long" else "ex-status"
    word_count = len(docx_text.split())
    st.markdown(
        f'<div class="{status_cls}">✅ <strong>{file_name}</strong> &nbsp;·&nbsp; '
        f'~{docx_est_pages} pages (estimated) &nbsp;·&nbsp; {word_count:,} words &nbsp;·&nbsp; '
        f'{len(docx_text):,} chars</div>',
        unsafe_allow_html=True)
    st.info("📝 Word document — using Single Document Mode. "
            "Book Mode is available for PDF files only.")

    if doc_size == "long":
        st.info("📋 **Long document** — multi-section summarisation will be used for full coverage.")
    elif doc_size == "medium":
        st.info("📋 **Medium document** — summarising in sections for good coverage.")

    # Set variables expected by Steps 4–7
    text          = docx_text
    sp, ep        = 1, docx_est_pages
    section_title = ""
    section_label = ""
    start_page, end_page = 1, docx_est_pages

    ga_event("pdf_processed",
             params={"file_name": file_name, "pages": f"1-{docx_est_pages}",
                     "mode": "docx", "doc_size": doc_size},
             once_key=f"processed_{file_name}_docx")

    # Skip Steps 2 and 3 entirely for Word files
    st.markdown('<div class="ex-section-label">Step 3 — Study Summary</div>', unsafe_allow_html=True)

else:
    # ── PDF path ──
    pdf_name    = file_name
    total_pages = get_pdf_page_count(file_bytes)

    # ============================================================
    # STEP 2 — Auto-detect document type (PDF only)
    # ============================================================
    is_likely_book = total_pages >= BOOK_PAGE_THRESHOLD

if st.session_state.doc_mode is None:
    st.markdown('<div class="ex-section-label">Step 2 — Choose study mode</div>', unsafe_allow_html=True)
    if is_likely_book:
        st.info(
            f"**{pdf_name}** has {total_pages} pages — looks like a book. "
            "Examora recommends **Book Mode** for chapter-by-chapter study.")
    else:
        st.info(
            f"**{pdf_name}** has {total_pages} pages — looks like a report or document. "
            "Examora recommends **Single Document Mode**.")
    ca, cb = st.columns(2)
    with ca:
        if st.button("📄 Single Document Mode", use_container_width=True, type="primary"):
            st.session_state.doc_mode = "single"; st.rerun()
    with cb:
        if st.button("📚 Book Mode (chapter-by-chapter)", use_container_width=True, type="primary"):
            st.session_state.doc_mode = "book"; st.rerun()
    st.stop()

doc_mode = st.session_state.doc_mode

with st.expander(f"Mode: **{'📄 Single Document' if doc_mode=='single' else '📚 Book'}** — click to switch"):
    sw1, sw2 = st.columns(2)
    with sw1:
        if st.button("Switch to Single Document Mode"):
            st.session_state.doc_mode = "single"; reset_doc_state(); st.rerun()
    with sw2:
        if st.button("Switch to Book Mode"):
            st.session_state.doc_mode = "book"; reset_doc_state(); st.rerun()


# ============================================================
# STEP 3A — Single Document Mode
# ============================================================
start_page, end_page = 1, total_pages
section_label = ""
section_title = ""
sp, ep = 1, total_pages

if doc_mode == "single":
    st.session_state.scope_mode = "Single Document"
    st.markdown(f'<div class="ex-section-label">📄 {pdf_name}</div>', unsafe_allow_html=True)

    scope = st.radio("Reading scope",
                     ["Whole document (recommended)", "Selected page range"],
                     horizontal=True, index=0)
    if scope.startswith("Whole"):
        start_page, end_page = 1, total_pages
    else:
        colA, colB = st.columns(2)
        with colA:
            start_page = st.number_input("Start page", min_value=1,
                                         max_value=max(1,total_pages), value=1, step=1)
        with colB:
            end_page = st.number_input("End page", min_value=1,
                                       max_value=max(1,total_pages),
                                       value=total_pages, step=1)

    _ck = (len(file_bytes), int(start_page), int(end_page))
    if st.session_state.get("_last_extracted_key") == _ck:
        text, _, sp, ep = extract_text_cached(file_bytes, int(start_page), int(end_page))
    else:
        with st.spinner("Reading document..."):
            text, _, sp, ep = extract_text_cached(file_bytes, int(start_page), int(end_page))
        st.session_state["_last_extracted_key"] = _ck

    if len(text) < 200:
        st.error("⚠️ Very little text extracted. This PDF may be scanned/image-based and requires OCR.")
        st.stop()
    elif len(text) < 500:
        st.warning("Small amount of text extracted. Consider widening the page range.")

    doc_size = classify_doc_size(text)
    status_cls = "ex-status-long" if doc_size == "long" else "ex-status"
    st.markdown(
        f'<div class="{status_cls}">✅ <strong>{pdf_name}</strong> &nbsp;·&nbsp; '
        f'{text_coverage_info(text, sp, ep)}</div>',
        unsafe_allow_html=True)

    if doc_size == "long":
        st.info(
            "📋 **Long document detected** — Examora will use multi-section summarisation "
            "to cover the full document. Summary generation will take slightly longer "
            "but will be comprehensive.")
    elif doc_size == "medium":
        st.info(
            "📋 **Medium document detected** — Examora will summarise in sections "
            "to ensure good coverage.")

    ga_event("pdf_processed",
             params={"pdf_name": pdf_name, "pages": f"{sp}-{ep}",
                     "mode": "single", "doc_size": doc_size},
             once_key=f"processed_{pdf_name}_{sp}_{ep}_single")


# ============================================================
# STEP 3B — Book Mode
# ============================================================
else:
    st.session_state.scope_mode = "Book Mode"
    st.markdown(f'<div class="ex-section-label">📚 {pdf_name}</div>', unsafe_allow_html=True)

    if st.session_state.toc_chapters is None:
        with st.spinner("Detecting chapters from PDF..."):
            st.session_state.toc_chapters = extract_toc_from_pdf(file_bytes, total_pages)

    toc_chapters = st.session_state.toc_chapters
    toc_found    = bool(toc_chapters)

    source_options = (
        ["Auto-detected chapters (TOC)", "Manual chapter list", "Quick split (every N pages)"]
        if toc_found else
        ["Manual chapter list", "Quick split (every N pages)"]
    )
    chapter_source = st.radio("Chapter source", source_options,
                              horizontal=True, key="book_chapter_source_radio")
    chapters: list[dict] = []

    if chapter_source.startswith("Auto"):
        chapters = toc_chapters

    elif chapter_source.startswith("Manual"):
        if not toc_found:
            st.info("No TOC found in this PDF. Enter chapter page ranges below.")
        default_ex = (
            "Chapter 1 - Basic Radiation Physics: 1-30\n"
            "Chapter 2 - Radiation Dosimetry: 31-65\n"
            "Chapter 3 - Treatment Planning: 66-100\n"
        )
        raw = st.text_area("Chapters (format: Chapter Name: start-end)",
                           value=st.session_state.manual_chapters_text or default_ex,
                           height=140)
        st.session_state.manual_chapters_text = raw
        for line in (raw or "").splitlines():
            line = line.strip()
            if not line:
                continue
            m = re.match(r"^(.+?)[\:\-–]\s*(\d+)\s*[-–]\s*(\d+)\s*$", line)
            if not m:
                continue
            name = m.group(1).strip()
            a, b = int(m.group(2)), int(m.group(3))
            a = max(1, min(a, total_pages))
            b = max(1, min(b, total_pages))
            if b < a: a, b = b, a
            chapters.append({"title": name, "start": a, "end": b})
        if not chapters:
            st.error("No valid chapters parsed. Use format: Chapter Name: start-end")
            st.stop()

    else:
        pages_per = st.number_input("Pages per section", min_value=5, max_value=100,
                                    value=st.session_state.quick_pages_per, step=5)
        st.session_state.quick_pages_per = int(pages_per)
        s, ch = 1, 1
        while s <= total_pages:
            e = min(total_pages, s + int(pages_per) - 1)
            chapters.append({"title": f"Section {ch} (pp. {s}–{e})", "start": s, "end": e})
            s = e + 1; ch += 1

    if not chapters:
        st.stop()

    # Bookmark banner
    bm_idx = load_bookmark(user_email, pdf_name)
    if bm_idx is not None and bm_idx < len(chapters):
        bm_ch = chapters[bm_idx]
        st.markdown(
            f'<div class="bookmark-banner">📌 <strong>Resume:</strong> '
            f'Last studied <strong>{bm_ch["title"]}</strong> '
            f'(pp. {bm_ch["start"]}–{bm_ch["end"]})</div>',
            unsafe_allow_html=True)
        rb1, rb2 = st.columns([1, 4])
        with rb1:
            if st.button("▶ Resume"):
                st.session_state.book_chapter_idx = bm_idx
                reset_exam_state()
                st.session_state.summary_open = False
                st.session_state.summary_text = ""
                st.rerun()
        with rb2:
            if st.button("✕ Clear bookmark"):
                clear_bookmark(user_email, pdf_name); st.rerun()

    chapter_idx  = min(st.session_state.book_chapter_idx, len(chapters) - 1)
    labels       = [f"{i+1}. {c['title']} (pp. {c['start']}–{c['end']})"
                    for i, c in enumerate(chapters)]
    selected_idx = st.selectbox("Select chapter", list(range(len(labels))),
                                index=chapter_idx, format_func=lambda i: labels[i])

    if selected_idx != st.session_state.book_chapter_idx:
        st.session_state.book_chapter_idx = selected_idx
        save_bookmark(user_email, pdf_name, selected_idx)
        reset_exam_state()
        st.session_state.summary_open = False
        st.session_state.summary_text = ""
        st.rerun()

    chapter_idx   = selected_idx
    current_ch    = chapters[chapter_idx]
    start_page    = current_ch["start"]
    end_page      = current_ch["end"]
    section_title = current_ch["title"]
    section_label = f" — {section_title}"

    nav1, nav2 = st.columns(2)
    with nav1:
        if st.button("◀ Previous chapter", disabled=(chapter_idx <= 0)):
            ni = chapter_idx - 1
            st.session_state.book_chapter_idx = ni
            save_bookmark(user_email, pdf_name, ni)
            reset_exam_state()
            st.session_state.summary_open = False
            st.session_state.summary_text = ""
            st.rerun()
    with nav2:
        if st.button("Next chapter ▶", disabled=(chapter_idx >= len(chapters)-1)):
            ni = chapter_idx + 1
            st.session_state.book_chapter_idx = ni
            save_bookmark(user_email, pdf_name, ni)
            reset_exam_state()
            st.session_state.summary_open = False
            st.session_state.summary_text = ""
            st.rerun()

    _ck = (len(file_bytes), int(start_page), int(end_page))
    if st.session_state.get("_last_extracted_key") == _ck:
        text, _, sp, ep = extract_text_cached(file_bytes, int(start_page), int(end_page))
    else:
        with st.spinner(f"Reading {section_title}..."):
            text, _, sp, ep = extract_text_cached(file_bytes, int(start_page), int(end_page))
        st.session_state["_last_extracted_key"] = _ck

    if len(text) < 200:
        st.error("⚠️ Very little text extracted. This chapter may be scanned/image-based.")
        st.stop()
    elif len(text) < 500:
        st.warning("Small amount of text extracted. Consider widening the chapter range.")

    st.markdown(
        f'<div class="ex-status">✅ <strong>{section_title}</strong> &nbsp;·&nbsp; '
        f'{text_coverage_info(text, sp, ep)} &nbsp;·&nbsp; '
        f'Chapter {chapter_idx+1} of {len(chapters)}</div>',
        unsafe_allow_html=True)

    ga_event("pdf_processed",
             params={"pdf_name": pdf_name, "pages": f"{sp}-{ep}", "mode": "book"},
             once_key=f"processed_{pdf_name}_{sp}_{ep}_book")


# ============================================================
# STEP 4 — Summary
# ============================================================
st.markdown('<div class="ex-section-label">Step 3 — Study Summary</div>', unsafe_allow_html=True)

doc_size = classify_doc_size(text)
n_chunks = 1 if doc_size == "short" else (3 if doc_size == "medium" else 5)
if n_chunks > 1:
    st.caption(
        f"ℹ️ This document will be summarised in {n_chunks} sections for complete coverage. "
        f"First generation takes ~{n_chunks * 8}–{n_chunks * 12} seconds.")

s_col1, s_col2, s_col3 = st.columns([1, 1, 2])
with s_col1:
    if st.button("📋 Generate Summary", type="primary", use_container_width=True):
        require_api_key()
        used = get_user_usage(user_email).get("exam_sessions_used", 0)
        if used >= BETA_LIMIT:
            st.error(f"Beta limit reached ({used}/{BETA_LIMIT}). Contact us for full access.")
            st.stop()

        th = get_text_hash(text)
        # Show section-by-section progress for multi-chunk docs
        if doc_size != "short":
            progress_bar = st.progress(0, text="Starting summarisation...")
            # We can't hook into the cached function's internals, so show
            # a pulsing message then call — result comes from cache on repeat
            progress_bar.progress(0.15, text=f"Analysing document ({doc_size})...")
            summary = safe_call(summarize_text_cached, th, text)
            progress_bar.progress(1.0, text="Summary complete ✓")
        else:
            with st.spinner("Generating summary..."):
                summary = safe_call(summarize_text_cached, th, text)

        if summary:
            increment_exam_session(user_email, {
                "ts": now_ts(), "pages": f"{sp}-{ep}",
                "type": "summary", "pdf_name": pdf_name,
                "doc_size": doc_size,
            })
            st.session_state.summary_text = summary
            st.session_state.summary_open = True
            st.rerun()

with s_col2:
    if st.session_state.summary_open:
        if st.button("✕ Close Summary", use_container_width=True):
            st.session_state.summary_open = False; st.rerun()

with s_col3:
    st.link_button("💬 Submit Feedback",
                   build_prefilled_feedback_url(FEEDBACK_URL, user_email))

# ── Render summary ──
if st.session_state.summary_open and st.session_state.summary_text:
    summary_text = st.session_state.summary_text
    subtitle = f"{pdf_name} &nbsp;·&nbsp; Pages {sp}–{ep}" + \
               (f" &nbsp;·&nbsp; {section_title}" if section_title else "")

    # Parse sections
    sections_map = {
        "Document Overview":  ("📄", "Document Overview"),
        "Key Concepts":       ("🔑", "Key Concepts"),
        "Common Confusions":  ("⚠️", "Common Confusions"),
        "Exam Takeaways":     ("🎯", "Exam Takeaways"),
    }
    current_section, current_lines = None, []
    parsed: dict[str, list[str]]   = {}
    for line in summary_text.splitlines():
        stripped = line.strip()
        matched  = False
        for key in sections_map:
            if stripped.startswith("##") and key.lower() in stripped.lower():
                if current_section:
                    parsed[current_section] = current_lines
                current_section = key
                current_lines   = []
                matched = True
                break
        if not matched and current_section:
            current_lines.append(line)
    if current_section:
        parsed[current_section] = current_lines

    size_label = {"short": "Short document", "medium": "Medium document", "long": "Long document"}
    st.markdown(f"""
<div class="sum-wrap">
  <div class="sum-header">
    <div class="sum-header-title">📋 Study Summary</div>
    <div class="sum-header-sub">{subtitle} &nbsp;·&nbsp; {size_label.get(doc_size,'')}</div>
  </div>
""", unsafe_allow_html=True)

    if parsed:
        for key, (icon, label) in sections_map.items():
            if key not in parsed:
                continue
            content = "\n".join(parsed[key]).strip()
            if not content:
                continue
            st.markdown(f"""
  <div class="sum-section">
    <div class="sum-section-title">{icon}&nbsp; {label}</div>
  </div>
""", unsafe_allow_html=True)
            st.markdown(content)
    else:
        st.markdown(summary_text)

    st.markdown("""
  <div class="sum-footer"></div>
</div>
""", unsafe_allow_html=True)

    st.download_button("⬇️ Download Summary (.txt)",
                       data=summary_text.encode("utf-8"),
                       file_name="examora_summary.txt", mime="text/plain")

with st.expander("Can't access the feedback form? Leave it here."):
    st.caption("Stored locally in .examora/feedback_log.json")
    fb_note = st.text_area("Your feedback", height=70, key="fb_note")
    if st.button("Save feedback"):
        if not (fb_note or "").strip():
            st.warning("Please type feedback first.")
        else:
            append_feedback_log({"ts": now_ts(), "email": user_email, "note": fb_note.strip()})
            st.success("Saved. Thank you!")


# ============================================================
# STEP 5 — Exam controls
# ============================================================
st.markdown('<div class="ex-section-label">Step 4 — Board Exam Practice</div>', unsafe_allow_html=True)

with st.expander("⚙️ Exam settings", expanded=not st.session_state.exam_open):
    left, right = st.columns([2, 1])
    with left:
        n_questions  = st.slider("Number of questions", 5, 25, value=10, step=5)
        difficulty   = st.selectbox("Difficulty", ["Easy","Medium","Hard"], index=1)
        shuffle_opts = st.checkbox("Shuffle answer options", value=True)
    with right:
        mode = st.selectbox("Feedback mode",
                            ["Exam Mode (feedback on submit)",
                             "Study Mode (instant feedback)"], index=0)

g1, g2, g3 = st.columns([1, 1, 2])
with g1:
    if st.button("🎯 Generate Exam", type="primary", use_container_width=True):
        require_api_key()
        used = get_user_usage(user_email).get("exam_sessions_used", 0)
        if used >= BETA_LIMIT:
            st.error(f"Beta limit reached ({used}/{BETA_LIMIT}). Contact us for full access.")
            st.stop()
        reset_exam_state()
        with st.spinner("Generating your exam..."):
            mcq_set = safe_call(generate_mcqs_cached, get_text_hash(text),
                                n_questions, difficulty, text)
        if mcq_set and mcq_set.get("questions"):
            questions = mcq_set["questions"]
            if shuffle_opts:
                questions = [shuffle_question_options(q) for q in questions]
            st.session_state.questions = questions
            st.session_state.exam_open = True
            new_used = increment_exam_session(user_email, {
                "ts": now_ts(), "pages": f"{sp}-{ep}",
                "n_questions": int(n_questions), "difficulty": difficulty,
                "type": "exam", "scope": st.session_state.scope_mode,
                "section_title": section_title, "pdf_name": pdf_name,
            })
            ga_event("exam_generated",
                     params={"pdf_name": pdf_name, "pages": f"{sp}-{ep}",
                             "n_questions": int(n_questions), "difficulty": difficulty},
                     once_key=f"exam_{pdf_name}_{sp}_{ep}_{n_questions}_{difficulty}")
            st.rerun()
        else:
            st.warning("No questions returned. Try a wider page range or lower difficulty.")

with g2:
    if st.button("✕ Close Exam", use_container_width=True):
        reset_exam_state(); st.rerun()

with g3:
    st.link_button("💬 Submit Feedback",
                   build_prefilled_feedback_url(FEEDBACK_URL, user_email))


# ============================================================
# STEP 6 — Exam panel
# ============================================================
if not st.session_state.exam_open or not st.session_state.questions:
    st.info("Click **Generate Exam** above to start your practice session.")
    st.stop()

questions      = st.session_state.questions
total          = len(questions)
study_mode_now = mode.startswith("Study Mode")
answered_count = sum(1 for i in range(1, total+1)
                     if st.session_state.answers.get(i) in ["A","B","C","D"])
subtitle_exam  = f"{pdf_name} &nbsp;·&nbsp; Pages {sp}–{ep}" + \
                 (f" &nbsp;·&nbsp; {section_title}" if section_title else "")

st.markdown(f"""
<div class="exam-wrap">
  <div class="exam-header">
    <div class="exam-header-title">🎯 Board Exam Practice</div>
    <div class="exam-header-sub">{subtitle_exam} &nbsp;·&nbsp; {difficulty} &nbsp;·&nbsp; {total} questions</div>
  </div>
</div>
""", unsafe_allow_html=True)

st.progress(answered_count / total if total else 0,
            text=f"Progress: {answered_count} / {total} questions answered")

st.markdown("**Jump to question:**")
grid_cols = st.columns(min(total, 10))
for i in range(total):
    qn         = i + 1
    ans_ok     = st.session_state.answers.get(qn) in ["A","B","C","D"]
    is_flagged = qn in st.session_state.flagged
    lbl        = str(qn) + (" ⚑" if is_flagged else "") + (" ✓" if ans_ok else "")
    if grid_cols[i % 10].button(lbl, key=f"grid_{qn}", disabled=st.session_state.submitted):
        st.session_state.q_index = i; st.rerun()

st.write("")

idx  = st.session_state.q_index
q    = questions[idx]
qnum = idx + 1
opts = q["options"]

st.markdown(f"""
<div class="q-card">
  <div class="q-meta">Question {qnum} of {total} &nbsp;·&nbsp; {difficulty}</div>
  <div class="q-text">{q['q']}</div>
</div>
""", unsafe_allow_html=True)

prev  = st.session_state.answers.get(qnum)
idx_r = ["A","B","C","D"].index(prev) if prev in ["A","B","C","D"] else None
choice = st.radio(
    "Select your answer", ["A","B","C","D"],
    format_func=lambda k: f"**{k})**  {opts[k]}",
    index=idx_r,
    key=f"choice_{qnum}",
    disabled=st.session_state.submitted,
    label_visibility="collapsed",
)
st.session_state.answers[qnum] = choice

if study_mode_now and not st.session_state.submitted and choice in ["A","B","C","D"]:
    if choice == q["answer"]:
        st.success("✅ Correct!")
    else:
        st.error(f"❌ Incorrect — Correct answer: **{q['answer']}**)  {opts[q['answer']]}")
    if q.get("explanation"):
        st.info(f"💡 {q['explanation']}")

n1, n2, n3, n4, n5 = st.columns([1, 1, 1, 2, 1])
with n1:
    if st.button("⚑ Flag", disabled=st.session_state.submitted):
        if qnum in st.session_state.flagged:
            st.session_state.flagged.remove(qnum)
        else:
            st.session_state.flagged.add(qnum)
        st.rerun()
with n2:
    if st.button("◀ Prev", disabled=st.session_state.submitted or idx == 0):
        st.session_state.q_index = idx - 1; st.rerun()
with n3:
    if st.button("Next ▶", disabled=st.session_state.submitted or idx == total-1):
        st.session_state.q_index = idx + 1; st.rerun()
with n4:
    submit_clicked = st.button("✅ Submit Exam", type="primary",
                               disabled=st.session_state.submitted, use_container_width=True)
with n5:
    st.write("")

if st.session_state.flagged:
    st.caption(f"⚑ Flagged: {', '.join(map(str, sorted(st.session_state.flagged)))}")

if submit_clicked:
    st.session_state.submitted = True
    ga_event("exam_submitted",
             params={"pdf_name": pdf_name, "pages": f"{sp}-{ep}",
                     "n_questions": total},
             once_key=f"submitted_{pdf_name}_{sp}_{ep}_{total}")
    st.rerun()


# ============================================================
# STEP 7 — Results
# ============================================================
if st.session_state.submitted:
    st.divider()
    correct  = sum(1 for i, qq in enumerate(questions, 1)
                   if st.session_state.answers.get(i) == qq["answer"])
    answered = sum(1 for i in range(1, total+1)
                   if st.session_state.answers.get(i) in ["A","B","C","D"])
    pct = int(correct / total * 100) if total else 0

    if pct >= 80:
        banner_cls, msg = "score-pass", "Excellent — ready to advance to the next topic."
    elif pct >= 60:
        banner_cls, msg = "score-warn", "Good attempt — review the explanations before moving on."
    else:
        banner_cls, msg = "score-fail", "More review needed — re-read this section and retry."

    st.markdown(f"""
<div class="{banner_cls}">
  <div class="score-pct">{pct}%</div>
  <div class="score-msg">
    {correct}/{total} correct &nbsp;·&nbsp; {answered}/{total} answered &nbsp;·&nbsp; {msg}
  </div>
</div>
""", unsafe_allow_html=True)

    st.markdown('<div class="ex-section-label">Question Review</div>', unsafe_allow_html=True)
    for i, qq in enumerate(questions, 1):
        user_ans    = st.session_state.answers.get(i)
        correct_ans = qq["answer"]
        is_correct  = user_ans == correct_ans
        css_cls     = "r-correct" if is_correct else "r-incorrect"
        icon        = "✅" if is_correct else "❌"
        wrong_note  = (f" &nbsp;·&nbsp; Correct: <strong>{correct_ans}</strong>) "
                       f"{qq['options'][correct_ans]}" if not is_correct else "")
        exp_html    = (f'<div class="r-exp">💡 {qq["explanation"]}</div>'
                       if qq.get("explanation") else "")
        st.markdown(f"""
<div class="{css_cls}">
  <div class="r-q">{icon} Q{i}: {qq['q']}</div>
  <div class="r-ans">Your answer: <strong>{user_ans or "No answer"}</strong>{wrong_note}</div>
  {exp_html}
</div>
""", unsafe_allow_html=True)

    st.divider()
    csv_bytes = build_results_csv(questions, st.session_state.answers)

    def _dl_evt():
        ga_event("results_downloaded",
                 params={"pdf_name": pdf_name, "pages": f"{sp}-{ep}", "n_questions": total},
                 once_key=f"dl_{pdf_name}_{sp}_{ep}_{total}")

    rc1, rc2, rc3 = st.columns([1, 1, 2])
    with rc1:
        st.download_button("⬇️ Download Results (CSV)", data=csv_bytes,
                           file_name="examora_results.csv",
                           mime="text/csv", on_click=_dl_evt)
    with rc2:
        if st.button("🔄 New Exam"):
            reset_exam_state(); st.rerun()

    st.divider()
    st.markdown("**Help us improve Examora** — report incorrect answers, unclear questions, or any issues.")
    st.link_button("💬 Submit Beta Feedback",
                   build_prefilled_feedback_url(FEEDBACK_URL, user_email))
